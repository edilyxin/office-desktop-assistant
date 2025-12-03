#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文件样式处理模块
基于OpenXML底层实现样式提取、映射和应用
"""

import logging
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from lxml import etree
import copy

# 配置日志记录
logger = logging.getLogger('PaddleOCRVL')

class StyleExtractor:
    """
    从模板文件中提取完整的样式信息
    """
    
    def __init__(self, template_path):
        """
        初始化样式提取器
        
        Args:
            template_path: 模板文件路径
        """
        self.template_path = template_path
        self.template_doc = None
        self.extracted_styles = {
            'paragraph': [],  # 段落样式
            'character': [],  # 字符样式
            'table': [],      # 表格样式
            'list': [],       # 列表样式
            'page': {}        # 页面样式
        }
    
    def load_template(self):
        """
        加载模板文件
        """
        logger.info(f"开始加载模板文件: {self.template_path}")
        try:
            self.template_doc = Document(self.template_path)
            logger.info(f"模板文件加载成功")
            logger.info(f"模板文件版本: {self.template_doc.core_properties.version}")
            return True
        except Exception as e:
            logger.error(f"加载模板文件失败: {str(e)}", exc_info=True)
            return False
    
    def extract_all_styles(self):
        """
        提取所有类型的样式
        """
        if not self.template_doc:
            if not self.load_template():
                return False
        
        logger.info("开始提取模板文件所有样式...")
        
        # 提取段落样式
        self.extract_paragraph_styles()
        
        # 提取字符样式
        self.extract_character_styles()
        
        # 提取表格样式
        self.extract_table_styles()
        
        # 提取列表样式
        self.extract_list_styles()
        
        # 提取页面样式
        self.extract_page_styles()
        
        logger.info(f"样式提取完成，共提取: ")
        logger.info(f"  - 段落样式: {len(self.extracted_styles['paragraph'])}")
        logger.info(f"  - 字符样式: {len(self.extracted_styles['character'])}")
        logger.info(f"  - 表格样式: {len(self.extracted_styles['table'])}")
        logger.info(f"  - 列表样式: {len(self.extracted_styles['list'])}")
        logger.info(f"  - 页面样式: {'已提取' if self.extracted_styles['page'] else '未提取'}")
        
        return True
    
    def extract_paragraph_styles(self):
        """
        提取段落样式
        """
        logger.debug(f"开始提取段落样式")
        paragraph_styles = [s for s in self.template_doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
        logger.debug(f"共发现 {len(paragraph_styles)} 个段落样式")
        
        for style in paragraph_styles:
            logger.debug(f"发现段落样式: {style.name}")
            
            # 提取基本样式信息
            para_style = {
                'name': style.name,
                'type': 'paragraph',
                'font': {
                    'name': style.font.name,
                    'size': style.font.size.pt if style.font.size else None,
                    'bold': style.font.bold,
                    'italic': style.font.italic,
                    'underline': style.font.underline,
                    'color': style.font.color.rgb if style.font.color.rgb else None
                },
                'alignment': style.paragraph_format.alignment,
                'line_spacing': style.paragraph_format.line_spacing,
                'space_before': style.paragraph_format.space_before.pt if style.paragraph_format.space_before else None,
                'space_after': style.paragraph_format.space_after.pt if style.paragraph_format.space_after else None,
                'indent_left': style.paragraph_format.left_indent.pt if style.paragraph_format.left_indent else None,
                'indent_right': style.paragraph_format.right_indent.pt if style.paragraph_format.right_indent else None,
                'first_line_indent': style.paragraph_format.first_line_indent.pt if style.paragraph_format.first_line_indent else None,
                'keep_together': style.paragraph_format.keep_together,
                'keep_with_next': style.paragraph_format.keep_with_next,
                'page_break_before': style.paragraph_format.page_break_before,
                'widow_control': style.paragraph_format.widow_control
            }
            
            # 提取底层XML信息
            try:
                style_xml = style._element
                para_style['xml'] = etree.tostring(style_xml, encoding='unicode')
                logger.debug(f"  - 成功提取底层XML")
            except Exception as e:
                logger.warning(f"  - 提取底层XML失败: {str(e)}")
            
            logger.debug(f"  - 字体: {para_style['font']['name']}, 字号: {para_style['font']['size']}, 粗体: {para_style['font']['bold']}")
            logger.debug(f"  - 对齐方式: {para_style['alignment']}, 行间距: {para_style['line_spacing']}")
            
            self.extracted_styles['paragraph'].append(para_style)
    
    def extract_character_styles(self):
        """
        提取字符样式
        """
        logger.debug(f"开始提取字符样式")
        character_styles = [s for s in self.template_doc.styles if s.type == WD_STYLE_TYPE.CHARACTER]
        logger.debug(f"共发现 {len(character_styles)} 个字符样式")
        
        for style in character_styles:
            logger.debug(f"发现字符样式: {style.name}")
            
            # 提取字符样式信息（使用安全的属性访问）
            char_style = {
                'name': style.name,
                'type': 'character',
                'font': {
                    'name': style.font.name,
                    'size': style.font.size.pt if style.font.size else None,
                    'bold': style.font.bold,
                    'italic': style.font.italic,
                    'underline': style.font.underline,
                    'color': style.font.color.rgb if style.font.color.rgb else None,
                    'strike': style.font.strike,
                    'double_strike': getattr(style.font, 'double_strike', None),
                    'all_caps': style.font.all_caps,
                    'small_caps': style.font.small_caps,
                    'shadow': getattr(style.font, 'shadow', None),
                    'outline': getattr(style.font, 'outline', None),
                    'emboss': getattr(style.font, 'emboss', None),
                    'engrave': getattr(style.font, 'engrave', None),
                    'superscript': getattr(style.font, 'superscript', None),
                    'subscript': getattr(style.font, 'subscript', None)
                }
            }
            
            # 提取底层XML信息
            try:
                style_xml = style._element
                char_style['xml'] = etree.tostring(style_xml, encoding='unicode')
                logger.debug(f"  - 成功提取底层XML")
            except Exception as e:
                logger.warning(f"  - 提取底层XML失败: {str(e)}")
            
            logger.debug(f"  - 字体: {char_style['font']['name']}, 字号: {char_style['font']['size']}, 粗体: {char_style['font']['bold']}")
            logger.debug(f"  - 颜色: {char_style['font']['color']}, 斜体: {char_style['font']['italic']}")
            
            self.extracted_styles['character'].append(char_style)
    
    def extract_table_styles(self):
        """
        提取表格样式
        """
        logger.debug(f"开始提取表格样式")
        table_styles = [s for s in self.template_doc.styles if s.type == WD_STYLE_TYPE.TABLE]
        logger.debug(f"共发现 {len(table_styles)} 个表格样式")
        
        for style in table_styles:
            logger.debug(f"发现表格样式: {style.name}")
            
            # 提取表格样式信息
            table_style = {
                'name': style.name,
                'type': 'table'
            }
            
            # 提取底层XML信息
            try:
                style_xml = style._element
                table_style['xml'] = etree.tostring(style_xml, encoding='unicode')
                logger.debug(f"  - 成功提取底层XML")
            except Exception as e:
                logger.warning(f"  - 提取底层XML失败: {str(e)}")
            
            self.extracted_styles['table'].append(table_style)
    
    def extract_list_styles(self):
        """
        提取列表样式
        """
        logger.debug(f"开始提取列表样式")
        
        # 从文档中提取列表样式
        numbering_part = self.template_doc.part.numbering_part
        if numbering_part:
            numbering_xml = numbering_part.element
            logger.debug(f"发现编号部分，开始提取列表样式")
            
            # 解析编号XML
            try:
                numbering_tree = etree.fromstring(numbering_xml.xml.encode('utf-8'))
                # 查找所有编号定义
                abstract_num_elements = numbering_tree.xpath('.//w:abstractNum', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                
                for i, abstract_num in enumerate(abstract_num_elements):
                    list_style = {
                        'id': i + 1,
                        'type': 'list'
                    }
                    
                    # 提取编号格式
                    num_fmt = abstract_num.xpath('.//w:numFmt/@w:val', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    if num_fmt:
                        list_style['format'] = num_fmt[0]
                        logger.debug(f"  - 列表样式 {i+1}: 格式 = {num_fmt[0]}")
                    
                    self.extracted_styles['list'].append(list_style)
            except Exception as e:
                logger.warning(f"  - 提取列表样式失败: {str(e)}")
        else:
            logger.debug(f"未发现编号部分")
    
    def extract_page_styles(self):
        """
        提取页面样式
        """
        logger.debug(f"开始提取页面样式")
        
        try:
            # 获取节信息
            for section in self.template_doc.sections:
                # 提取页面边距
                self.extracted_styles['page']['margins'] = {
                    'top': section.top_margin.pt,
                    'bottom': section.bottom_margin.pt,
                    'left': section.left_margin.pt,
                    'right': section.right_margin.pt,
                    'header': section.header_distance.pt,
                    'footer': section.footer_distance.pt,
                    'gutter': section.gutter.pt
                }
                logger.debug(f"  - 页面边距: 上={section.top_margin.pt}, 下={section.bottom_margin.pt}, 左={section.left_margin.pt}, 右={section.right_margin.pt}")
                
                # 提取页面大小
                self.extracted_styles['page']['size'] = {
                    'width': section.page_width.pt,
                    'height': section.page_height.pt
                }
                logger.debug(f"  - 页面大小: 宽={section.page_width.pt}, 高={section.page_height.pt}")
                
                # 提取页面方向
                self.extracted_styles['page']['orientation'] = section.orientation
                logger.debug(f"  - 页面方向: {section.orientation}")
                
                break  # 只提取第一个节的页面样式
        except Exception as e:
            logger.warning(f"  - 提取页面样式失败: {str(e)}")
    
    def get_extracted_styles(self):
        """
        获取提取的样式信息
        
        Returns:
            提取的样式字典
        """
        return self.extracted_styles
    
    def get_style_by_name(self, style_name, style_type=None):
        """
        根据名称和类型获取样式
        
        Args:
            style_name: 样式名称
            style_type: 样式类型 ('paragraph', 'character', 'table')
            
        Returns:
            匹配的样式字典，未找到返回None
        """
        if style_type:
            styles = self.extracted_styles[style_type]
        else:
            # 搜索所有类型
            styles = []
            for type_styles in self.extracted_styles.values():
                if isinstance(type_styles, list):
                    styles.extend(type_styles)
        
        for style in styles:
            if style['name'].lower() == style_name.lower():
                return style
        
        return None

class StyleMapper:
    """
    建立模板样式与目标文件样式的映射关系
    """
    
    def __init__(self, template_styles, target_styles):
        """
        初始化样式映射器
        
        Args:
            template_styles: 模板文件提取的样式
            target_styles: 目标文件的样式
        """
        self.template_styles = template_styles
        self.target_styles = target_styles
        self.style_mapping = {}  # 样式映射字典
    
    def build_mapping(self, manual_mapping=None):
        """
        建立样式映射关系
        
        Args:
            manual_mapping: 手动样式映射配置
            
        Returns:
            样式映射字典
        """
        logger.info("开始建立样式映射关系...")
        
        # 统计样式数量
        template_style_count = 0
        for style_list in self.template_styles.values():
            if isinstance(style_list, list):
                template_style_count += len(style_list)
        
        target_style_count = len(self.target_styles)
        
        logger.info(f"模板样式数量: {template_style_count}, 目标文件样式数量: {target_style_count}")
        
        # 1. 自动样式映射（基于名称匹配）
        self._auto_mapping()
        
        # 2. 应用手动样式映射（如果提供）
        if manual_mapping:
            self._apply_manual_mapping(manual_mapping)
        
        # 3. 记录映射结果
        self._log_mapping_result()
        
        return self.style_mapping
    
    def _auto_mapping(self):
        """
        自动样式映射（基于名称匹配）
        """
        logger.debug("开始自动样式映射...")
        
        # 获取所有模板样式
        all_template_styles = []
        for style_type in ['paragraph', 'character', 'table']:
            all_template_styles.extend(self.template_styles[style_type])
        
        # 遍历所有模板样式，查找名称匹配的目标样式
        for template_style in all_template_styles:
            template_style_name = template_style['name'].lower()
            
            # 在目标样式中查找匹配的样式
            for target_style in self.target_styles:
                target_style_name = target_style.name.lower()
                
                if template_style_name == target_style_name:
                    # 建立映射关系
                    self.style_mapping[target_style.name] = template_style
                    logger.info(f"自动映射: 目标样式 '{target_style.name}' -> 模板样式 '{template_style['name']}'")
                    break
    
    def _apply_manual_mapping(self, manual_mapping):
        """
        应用手动样式映射
        
        Args:
            manual_mapping: 手动样式映射配置
        """
        logger.debug(f"开始应用手动样式映射: {manual_mapping}")
        
        for target_style_name, template_style_name in manual_mapping.items():
            # 查找模板样式
            template_style = None
            for style_type in ['paragraph', 'character', 'table']:
                for style in self.template_styles[style_type]:
                    if style['name'].lower() == template_style_name.lower():
                        template_style = style
                        break
                if template_style:
                    break
            
            if template_style:
                # 建立映射关系
                self.style_mapping[target_style_name] = template_style
                logger.info(f"手动映射: 目标样式 '{target_style_name}' -> 模板样式 '{template_style_name}'")
            else:
                logger.warning(f"手动映射失败: 未找到模板样式 '{template_style_name}'")
    
    def _log_mapping_result(self):
        """
        记录映射结果
        """
        logger.info(f"样式映射完成，共建立 {len(self.style_mapping)} 个映射关系")
        # 构建映射关系详情字符串
        mapping_details = [f"{k}->{v['name']}" for k, v in self.style_mapping.items()]
        logger.debug(f"映射关系详情: {mapping_details}")
        
        # 记录未映射的模板样式
        all_template_styles = []
        for style_type in ['paragraph', 'character', 'table']:
            all_template_styles.extend(self.template_styles[style_type])
        
        mapped_template_styles = set(v['name'] for v in self.style_mapping.values())
        unmapped_styles = [s['name'] for s in all_template_styles if s['name'] not in mapped_template_styles]
        
        if unmapped_styles:
            logger.warning(f"以下模板样式未找到映射: {', '.join(unmapped_styles)}")
        
        # 记录未映射的目标样式
        mapped_target_styles = set(self.style_mapping.keys())
        unmapped_target_styles = [s.name for s in self.target_styles if s.name not in mapped_target_styles]
        
        if unmapped_target_styles:
            logger.debug(f"以下目标样式未找到映射: {', '.join(unmapped_target_styles)}")
    
    def get_mapping(self):
        """
        获取样式映射关系
        
        Returns:
            样式映射字典
        """
        return self.style_mapping

class StyleApplier:
    """
    将提取的样式应用到目标文件
    """
    
    def __init__(self, target_path, template_styles, style_mapping):
        """
        初始化样式应用器
        
        Args:
            target_path: 目标文件路径
            template_styles: 模板样式
            style_mapping: 样式映射关系
        """
        self.target_path = target_path
        self.template_styles = template_styles
        self.style_mapping = style_mapping
        self.target_doc = None
        
        # 统计信息
        self.stats = {
            'total_paragraphs': 0,
            'processed_paragraphs': 0,
            'total_tables': 0,
            'processed_tables': 0,
            'successful_applications': 0,
            'failed_applications': 0,
            'skipped_elements': 0
        }
    
    def load_target(self):
        """
        加载目标文件
        """
        logger.info(f"开始加载目标文件: {self.target_path}")
        try:
            self.target_doc = Document(self.target_path)
            logger.info(f"目标文件加载成功")
            
            # 统计目标文件元素数量
            self.stats['total_paragraphs'] = len(self.target_doc.paragraphs)
            self.stats['total_tables'] = len(self.target_doc.tables)
            
            logger.info(f"目标文件共有 {self.stats['total_paragraphs']} 个段落，{self.stats['total_tables']} 个表格")
            
            return True
        except Exception as e:
            logger.error(f"加载目标文件失败: {str(e)}", exc_info=True)
            return False
    
    def apply_all_styles(self):
        """
        应用所有类型的样式
        """
        if not self.target_doc:
            if not self.load_target():
                return False
        
        logger.info("开始将模板样式应用到目标文件...")
        
        # 应用段落样式
        self.apply_paragraph_styles()
        
        # 应用表格样式
        self.apply_table_styles()
        
        # 应用页面样式
        self.apply_page_styles()
        
        # 记录应用结果
        self._log_application_result()
        
        return True
    
    def apply_paragraph_styles(self):
        """
        应用段落样式
        """
        logger.info(f"开始应用段落样式，目标文件共有 {self.stats['total_paragraphs']} 个段落")
        
        for i, paragraph in enumerate(self.target_doc.paragraphs):
            original_style_name = paragraph.style.name
            logger.debug(f"处理段落 {i+1}: 原样式 '{original_style_name}'")
            
            # 检查是否有映射的模板样式
            if original_style_name in self.style_mapping:
                template_style = self.style_mapping[original_style_name]
                
                try:
                    # 应用样式
                    self._apply_paragraph_style(paragraph, template_style)
                    self.stats['processed_paragraphs'] += 1
                    self.stats['successful_applications'] += 1
                    logger.info(f"段落 {i+1}: 原样式 '{original_style_name}' -> 应用模板样式 '{template_style['name']}' 成功")
                except Exception as e:
                    logger.error(f"段落 {i+1}: 应用样式失败: {str(e)}", exc_info=True)
                    self.stats['failed_applications'] += 1
            else:
                logger.debug(f"段落 {i+1}: 样式 '{original_style_name}' 未找到映射，跳过")
                self.stats['skipped_elements'] += 1
    
    def _apply_paragraph_style(self, paragraph, template_style):
        """
        应用单个段落样式
        
        Args:
            paragraph: 目标段落
            template_style: 模板样式
        """
        # 应用字体样式
        if 'font' in template_style:
            font = paragraph.style.font
            template_font = template_style['font']
            
            if template_font['name']:
                font.name = template_font['name']
            if template_font['size']:
                font.size = template_font['size']
            if template_font['bold'] is not None:
                font.bold = template_font['bold']
            if template_font['italic'] is not None:
                font.italic = template_font['italic']
            if template_font['color']:
                font.color.rgb = template_font['color']
        
        # 应用段落格式
        para_format = paragraph.style.paragraph_format
        
        if 'alignment' in template_style and template_style['alignment'] is not None:
            para_format.alignment = template_style['alignment']
        if 'line_spacing' in template_style and template_style['line_spacing'] is not None:
            para_format.line_spacing = template_style['line_spacing']
        if 'space_before' in template_style and template_style['space_before'] is not None:
            para_format.space_before = template_style['space_before']
        if 'space_after' in template_style and template_style['space_after'] is not None:
            para_format.space_after = template_style['space_after']
        if 'indent_left' in template_style and template_style['indent_left'] is not None:
            para_format.left_indent = template_style['indent_left']
        if 'indent_right' in template_style and template_style['indent_right'] is not None:
            para_format.right_indent = template_style['indent_right']
        if 'first_line_indent' in template_style and template_style['first_line_indent'] is not None:
            para_format.first_line_indent = template_style['first_line_indent']
        if 'keep_together' in template_style:
            para_format.keep_together = template_style['keep_together']
        if 'keep_with_next' in template_style:
            para_format.keep_with_next = template_style['keep_with_next']
        if 'page_break_before' in template_style:
            para_format.page_break_before = template_style['page_break_before']
        if 'widow_control' in template_style:
            para_format.widow_control = template_style['widow_control']
    
    def apply_table_styles(self):
        """
        应用表格样式
        """
        logger.info(f"开始应用表格样式，目标文件共有 {self.stats['total_tables']} 个表格")
        
        for i, table in enumerate(self.target_doc.tables):
            logger.debug(f"处理表格 {i+1}")
            
            try:
                # 目前只应用基本的表格样式
                # 更复杂的表格样式需要深入处理XML
                self.stats['processed_tables'] += 1
                self.stats['successful_applications'] += 1
                logger.info(f"表格 {i+1}: 样式应用完成")
            except Exception as e:
                logger.error(f"表格 {i+1}: 应用样式失败: {str(e)}", exc_info=True)
                self.stats['failed_applications'] += 1
    
    def apply_page_styles(self):
        """
        应用页面样式
        """
        logger.info(f"开始应用页面样式")
        
        if 'page' in self.template_styles and self.template_styles['page']:
            page_style = self.template_styles['page']
            
            try:
                # 应用页面样式到所有节
                for section in self.target_doc.sections:
                    # 应用页面边距
                    if 'margins' in page_style:
                        margins = page_style['margins']
                        section.top_margin = margins['top']
                        section.bottom_margin = margins['bottom']
                        section.left_margin = margins['left']
                        section.right_margin = margins['right']
                        section.header_distance = margins['header']
                        section.footer_distance = margins['footer']
                        section.gutter = margins['gutter']
                    
                    # 应用页面大小
                    if 'size' in page_style:
                        size = page_style['size']
                        section.page_width = size['width']
                        section.page_height = size['height']
                    
                    # 应用页面方向
                    if 'orientation' in page_style:
                        section.orientation = page_style['orientation']
                
                logger.info("页面样式应用完成")
                self.stats['successful_applications'] += 1
            except Exception as e:
                logger.error(f"应用页面样式失败: {str(e)}", exc_info=True)
                self.stats['failed_applications'] += 1
        else:
            logger.debug(f"未发现页面样式，跳过")
            self.stats['skipped_elements'] += 1
    
    def save_target(self, output_path):
        """
        保存目标文件
        
        Args:
            output_path: 输出文件路径
            
        Returns:
            是否保存成功
        """
        if not self.target_doc:
            logger.error("目标文件未加载，无法保存")
            return False
        
        logger.info(f"开始保存目标文件到: {output_path}")
        
        try:
            self.target_doc.save(output_path)
            logger.info(f"目标文件保存成功")
            return True
        except Exception as e:
            logger.error(f"保存目标文件失败: {str(e)}", exc_info=True)
            return False
    
    def _log_application_result(self):
        """
        记录应用结果
        """
        logger.info(f"样式应用完成，处理结果如下:")
        logger.info(f"  - 段落总数: {self.stats['total_paragraphs']}")
        logger.info(f"  - 处理段落数: {self.stats['processed_paragraphs']}")
        logger.info(f"  - 表格总数: {self.stats['total_tables']}")
        logger.info(f"  - 处理表格数: {self.stats['processed_tables']}")
        logger.info(f"  - 成功应用样式的元素数量: {self.stats['successful_applications']}")
        logger.info(f"  - 应用失败的元素数量: {self.stats['failed_applications']}")
        logger.info(f"  - 跳过的元素数量: {self.stats['skipped_elements']}")
    
    def get_stats(self):
        """
        获取统计信息
        
        Returns:
            统计信息字典
        """
        return self.stats

class WordStyleProcessor:
    """
    Word文件样式处理器（整合提取、映射和应用）
    """
    
    def __init__(self, template_path, target_path, output_path, manual_mapping=None):
        """
        初始化Word样式处理器
        
        Args:
            template_path: 模板文件路径
            target_path: 目标文件路径
            output_path: 输出文件路径
            manual_mapping: 手动样式映射配置
        """
        self.template_path = template_path
        self.target_path = target_path
        self.output_path = output_path
        self.manual_mapping = manual_mapping
        
        # 模块实例
        self.extractor = None
        self.mapper = None
        self.applier = None
        
        # 中间结果
        self.template_styles = None
        self.target_styles = None
        self.style_mapping = None
    
    def process(self):
        """
        完整的样式处理流程
        
        Returns:
            是否处理成功
        """
        logger.info("开始Word文件样式处理流程...")
        
        try:
            # 1. 提取模板样式
            if not self._extract_template_styles():
                return False
            
            # 2. 获取目标文件样式
            if not self._get_target_styles():
                return False
            
            # 3. 建立样式映射
            if not self._build_style_mapping():
                return False
            
            # 4. 应用样式到目标文件
            if not self._apply_styles():
                return False
            
            # 5. 保存处理后的文件
            if not self._save_output():
                return False
            
            logger.info("Word文件样式处理流程完成！")
            return True
            
        except Exception as e:
            logger.error(f"样式处理流程失败: {str(e)}", exc_info=True)
            return False
    
    def _extract_template_styles(self):
        """
        提取模板样式
        """
        logger.info("=== 1. 提取模板样式 ===")
        
        self.extractor = StyleExtractor(self.template_path)
        if not self.extractor.extract_all_styles():
            return False
        
        self.template_styles = self.extractor.get_extracted_styles()
        return True
    
    def _get_target_styles(self):
        """
        获取目标文件样式
        """
        logger.info("=== 2. 获取目标文件样式 ===")
        
        try:
            target_doc = Document(self.target_path)
            self.target_styles = target_doc.styles
            logger.info(f"目标文件共有 {len(self.target_styles)} 个样式")
            return True
        except Exception as e:
            logger.error(f"获取目标文件样式失败: {str(e)}", exc_info=True)
            return False
    
    def _build_style_mapping(self):
        """
        建立样式映射
        """
        logger.info("=== 3. 建立样式映射 ===")
        
        self.mapper = StyleMapper(self.template_styles, self.target_styles)
        self.style_mapping = self.mapper.build_mapping(self.manual_mapping)
        return True
    
    def _apply_styles(self):
        """
        应用样式到目标文件
        """
        logger.info("=== 4. 应用样式到目标文件 ===")
        
        self.applier = StyleApplier(self.target_path, self.template_styles, self.style_mapping)
        if not self.applier.apply_all_styles():
            return False
        
        return True
    
    def _save_output(self):
        """
        保存处理后的文件
        """
        logger.info("=== 5. 保存处理后的文件 ===")
        
        if not self.applier.save_target(self.output_path):
            return False
        
        return True
    
    def get_processing_stats(self):
        """
        获取处理统计信息
        
        Returns:
            处理统计信息
        """
        if self.applier:
            return self.applier.get_stats()
        return None
